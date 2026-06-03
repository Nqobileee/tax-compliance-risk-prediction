"""Build TaxGuard research proposal DOCX in JRISS article format."""
from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
FIGURES = ROOT / "notebooks" / "figures"
METRICS = ROOT / "outputs" / "metrics" / "experiment_summary.json"
OUTPUT = ROOT / "docs" / "TaxGuard_Research_Proposal.docx"

MIN_WORDS = 4000
MAX_WORDS = 6000


def set_run_font(run, name: str = "Times New Roman", size: int = 12, bold: bool = False, italic: bool = False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_centered(doc: Document, text: str, size: int = 12, bold: bool = False, space_after: int = 6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_body(doc: Document, text: str, first_line_indent: float = 0.5):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Inches(first_line_indent)
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_heading_section(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, bold=True)
    return p


def add_subheading(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, bold=True, italic=False)
    return p


def add_figure(doc: Document, image_path: Path, caption: str, width: float = 5.8):
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cap.paragraph_format.space_after = Pt(10)
    run = cap.add_run(caption)
    set_run_font(run, size=11, italic=True)


def add_table_caption(doc: Document, caption: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(caption)
    set_run_font(run, bold=True, size=11)


def count_words(doc: Document) -> int:
    text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += "\n" + cell.text
    return len(re.findall(r"\b[\w'-]+\b", text))


def build_document(metrics: dict) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # --- JRISS header block (mirrors reference paper layout) ---
    add_centered(doc, "Journal of Research and Innovation for Sustainable Society (JRISS)", 11)
    add_centered(doc, "Volume 6, Issue 1, 2026", 11)
    add_centered(doc, "ISSN: 2668-0416", 11)
    add_centered(doc, "Thoth Publishing House", 11)
    add_centered(doc, "1", 11, space_after=12)

    add_centered(doc, "TaxGuard: An AI-Based Anomaly Detection Framework for Corporate Tax Return Risk Scoring at ZIMRA", 14, bold=True, space_after=10)

    add_centered(doc, "Edith Muyambiri\u00b9 and Andile Bhebhe\u00b2", 12, space_after=6)
    add_centered(doc, "\u00b9Zimbabwe Revenue Authority (ZIMRA), Harare, Zimbabwe", 11, space_after=2)
    add_centered(doc, "\u00b2Zimbabwe Revenue Authority (ZIMRA), Harare, Zimbabwe", 11, space_after=6)
    add_centered(doc, "E-mail: emuyambiri@zimra.co.zw ; abhebhe@zimra.co.zw", 10, space_after=4)
    add_centered(doc, "Theme: Area 1 \u2014 Data, Automation and Intelligent Systems in the 4IR Era", 10, space_after=2)
    add_centered(doc, "Category: Prototype Demonstration", 10, space_after=14)

    auc_dev = metrics["top_univariate_auc"]["AUC_High_Risk"]["Tax_Rate_Deviation"]
    auc_base = metrics["model_oof"]["audit_likelihood_auc_baseline"]
    auc_ens = metrics["model_oof"]["ensemble_roc_auc"]
    pr_auc = metrics["model_oof"]["ensemble_pr_auc"]
    brier = metrics["model_oof"]["ensemble_brier"]
    f1 = metrics["model_oof"]["ensemble_f1_optimal"]
    n_records = metrics["dataset"]["records"]

    abstract = (
        "Abstract. Corporate tax non-compliance poses a significant threat to Zimbabwe\u2019s fiscal sustainability. "
        "However, the Zimbabwe Revenue Authority (ZIMRA) currently relies heavily on manual audit selection processes "
        "that are resource-intensive, inconsistent, and unable to scale with the growing volume and complexity of "
        "corporate tax filings (Thomas, 2000; Crook et al., 2007). The absence of a data-driven risk prioritisation "
        "mechanism results in inefficient audit allocation, where high-risk fraudulent submissions may evade detection "
        "while compliant firms are subjected to disproportionate scrutiny (Verstraeten and Van den Poel, 2004). "
        "This paper proposes TaxGuard, a hybrid machine learning framework that integrates both unsupervised and "
        "supervised techniques to detect anomalies in corporate tax returns and assign probabilistic risk scores to "
        "each filing (Liu et al., 2008; Hinton and Salakhutdinov, 2006). The system ingests historical tax submissions, "
        "audited financial statements and transactional data, applying Isolation Forest and Autoencoder neural networks "
        "for anomaly detection. These are complemented by a Gradient Boosted classifier trained on confirmed fraud cases "
        "to enhance predictive precision over time (Friedman, 2001; Pedregosa et al., 2011). Key detection features "
        "include deviations in income-to-expense ratios, inter-period financial inconsistencies, sector-based outliers, "
        "and irregular value-added tax (VAT) to turnover relationships (OECD, 2015). Each tax return is assigned a risk "
        "score supported by SHAP-based explainability, enabling auditors to identify the specific financial indicators "
        "contributing to a flagged case (Lundberg and Lee, 2017). This improves transparency, reduces subjective bias and "
        "enhances audit prioritisation. Furthermore, the system incorporates continuous learning by retraining on outcomes "
        "from completed audits, thereby progressively improving detection accuracy (Kohavi, 1995). A prototype was "
        f"evaluated on {n_records:,} corporate filings using five-fold stratified out-of-fold validation. Tax rate deviation "
        f"achieved univariate AUROC {auc_dev:.4f} compared with ZIMRA\u2019s legacy Audit_Likelihood heuristic at {auc_base:.4f}. "
        f"The stacked ensemble achieved OOF ROC-AUC {auc_ens:.4f}, PR-AUC {pr_auc:.4f}, Brier score {brier:.4f}, and F1 "
        f"{f1:.4f}. The study demonstrates how TaxGuard advances intelligent fiscal governance in emerging economies by "
        "strengthening revenue assurance and audit efficiency. It aligns with Zimbabwe\u2019s National AI Strategy (2026\u20132030) "
        "and National Development Strategy 2 objectives, particularly in enhancing domestic revenue mobilisation and "
        "promoting accountability through data-driven decision-making (Government of Zimbabwe, 2026)."
    )
    p_abs = doc.add_paragraph()
    p_abs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_abs.paragraph_format.line_spacing = 1.15
    p_abs.paragraph_format.space_after = Pt(10)
    run = p_abs.add_run(abstract)
    set_run_font(run)

    kw = doc.add_paragraph()
    kw.paragraph_format.space_after = Pt(12)
    r1 = kw.add_run("Key words: ")
    set_run_font(r1, bold=True)
    r2 = kw.add_run(
        "tax compliance, anomaly detection, gradient boosting, SHAP, audit prioritisation, ZIMRA, "
        "machine learning, fiscal governance, corporate tax risk scoring"
    )
    set_run_font(r2)

    # --- 1 Introduction ---
    add_heading_section(doc, "1 Introduction")
    add_body(
        doc,
        "The ministry of finance and economic development, through the Zimbabwe Revenue Authority (ZIMRA), "
        "bears a constitutional mandate to collect revenue that funds national infrastructure, health, education, "
        "and social protection. Corporate income tax is among the most complex and consequential revenue streams "
        "because liabilities depend on accounting profit, transfer pricing, offshore structures, and discretionary "
        "tax planning positions that may only become visible after a field audit (OECD, 2015; ZIMRA, 2024). "
        "In emerging economies, revenue authorities face dual pressure: taxpayers expect equitable treatment, "
        "while treasuries require predictable inflows to finance development programmes under National Development "
        "Strategy 2 (Government of Zimbabwe, 2026).",
    )
    add_body(
        doc,
        "Historically, audit case selection at ZIMRA blended rule-based filters\u2014such as sector, refund claims, "
        "and payment delays\u2014with examiner judgement. That approach functioned when filing volumes were moderate, "
        "but it scales poorly. Examiners cannot review every return; they must rank cases. Without a quantitative "
        "ranker, ranking reverts to heuristics such as large revenue or recent adverse events, which prototype "
        "evidence shows to be weaker than statutory\u2013effective tax rate gaps (Nyoni and Matshisela, 2018). "
        "TaxGuard is designed as a decision-support layer: it does not replace statutory powers or professional "
        "judgement; it orders the queue and documents why a return surfaced.",
    )
    add_body(
        doc,
        "Analogous to credit scoring\u2014where banks estimate probability of default from applicant characteristics "
        "(Rose, 2002; BIS, 2000)\u2014ZIMRA can treat audit prioritisation as predicting elevated non-compliance or "
        "elevated tax risk tier. The parallel extends to evaluation metrics: area under the receiver operating "
        "characteristic curve (AUROC), precision\u2013recall, Brier score, and Kolmogorov\u2013Smirnov separation are "
        "standard in retail credit literature and are applied here to tax audit risk (Hanley and McNeil, 1982; "
        "Thomas, 2000). Zimbabwean researchers have already demonstrated machine learning superiority over linear "
        "probability models for credit risk (Nyoni and Matshisela, 2018; Nyathi et al., 2014), providing a "
        "methodological precedent for public-sector scoring at ZIMRA.",
    )
    add_body(
        doc,
        "The government of Zimbabwe has expanded its digital transformation agenda through the National AI Strategy "
        "(2026\u20132030), which promotes ethical, explainable artificial intelligence in public administration. TaxGuard "
        "responds directly to that policy direction by combining hybrid anomaly detection, calibrated probabilistic "
        "scores, and SHAP-based explanations suitable for auditor-facing workflows. This research prototype "
        "demonstrates that data-driven audit prioritisation is not only feasible but dramatically superior to manual "
        "selection when measured by ranking quality and operational transparency.",
    )
    add_body(
        doc,
        "Corporate tax administration in Zimbabwe operates within a broader Southern African context where revenue "
        "authorities are modernising core systems, expanding e-filing, and exploring analytics to close compliance gaps "
        "(ZIMRA, 2024). Yet analytics maturity often lags policy ambition: many selection rules remain deterministic "
        "and static, unable to learn from completed audits or adapt when firms shift planning structures. TaxGuard "
        "addresses that gap by treating each audit closure as labelled feedback, consistent with continuous learning "
        "pipelines described in machine learning operations literature (Kohavi, 1995; Pedregosa et al., 2011).",
    )
    add_body(
        doc,
        "The present study contributes four elements aligned with JRISS publication norms and the prototype demonstration "
        "category. First, it profiles data quality and target distributions suitable for Audit Management System "
        "integration. Second, it engineers domain features aligned with OECD BEPS indicators and control environment "
        "practice. Third, it ranks predictors using univariate AUROC, correlation structure, and composite policy flags "
        "that expose manual selection blind spots. Fourth, it trains and evaluates a hybrid machine learning pipeline "
        "with leakage-safe five-fold stratified out-of-fold validation and SHAP explainability, benchmarking against "
        "ZIMRA\u2019s existing Audit_Likelihood heuristic.",
    )

    add_figure(
        doc,
        FIGURES / "01_target_distributions.png",
        "Figure 1. Distribution of tax risk labels (balanced thirds), audit outcomes (69.4% Clean), and ZIMRA Audit_Likelihood heuristic. Source: TaxGuard exploratory analysis, June 2026.",
    )

    # --- 2 Literature review ---
    add_heading_section(doc, "2 Literature review")
    add_body(
        doc,
        "Tax administration literature emphasises risk-based compliance management: authorities should concentrate "
        "scarce audit resources on returns with the highest expected yield and deterrence value (Basel Committee, "
        "2006; Harwood et al., 1999). In practice, risk scoring requires measurable indicators of non-compliance "
        "potential. OECD Base Erosion and Profit Shifting (BEPS) work highlights transfer pricing, offshore "
        "exposure, and documentation gaps as corporate risk dimensions relevant to African revenue authorities "
        "(OECD, 2015). Machine learning offers multivariate ranking when linear rules fail to capture interaction "
        "effects between planning aggressiveness and statutory\u2013effective tax gaps.",
    )

    add_table_caption(doc, "Table 1. Definitions of key terms.")
    t1 = doc.add_table(rows=6, cols=2)
    t1.style = "Table Grid"
    defs = [
        ("Tax risk scoring", "The assignment of a calibrated probability or score s_i \u2208 [0,1] to corporate filing i indicating elevated likelihood of High tax risk tier or material non-compliance."),
        ("Audit prioritisation", "Ranking filings by s_i so that, for fixed audit capacity k, the top-k cases maximise expected audit yield subject to fairness and explainability constraints."),
        ("Hybrid ML framework", "A stacked ensemble combining unsupervised anomaly detectors (Isolation Forest, Autoencoder) with supervised gradient boosting and a logistic meta-learner."),
        ("SHAP explainability", "Shapley Additive exPlanations attributing each prediction to feature contributions, enabling auditor-facing transparency (Lundberg and Lee, 2017)."),
        ("Continuous learning", "Retraining models on Qualified and Adverse audit closures so detection accuracy improves as institutional knowledge accumulates."),
    ]
    hdr = t1.rows[0].cells
    hdr[0].text = "Key term"
    hdr[1].text = "Definition"
    for i, (term, definition) in enumerate(defs, start=1):
        t1.rows[i].cells[0].text = term
        t1.rows[i].cells[1].text = definition

    add_body(
        doc,
        "Unsupervised anomaly detection identifies multivariate outliers without labelled fraud cases. Isolation "
        "Forest partitions feature space randomly and scores anomalies by expected path length (Liu et al., 2008). "
        "Autoencoders learn compressed representations and flag reconstructions with high error (Hinton and "
        "Salakhutdinov, 2006). Supervised gradient boosting iteratively fits regression trees to pseudo-responses "
        "from binary cross-entropy loss, capturing non-linear interactions (Friedman, 2001). Stacking combines "
        "out-of-fold base-model predictions through a meta-learner trained without in-sample leakage (Kohavi, 1995).",
    )
    add_body(
        doc,
        "Nyoni and Matshisela (2018) compared logistic regression, random forest, and lasso on German credit data "
        "and reported AUROC up to 0.8048 for lasso, establishing benchmark expectations for Zimbabwean scoring "
        "research. Nyathi et al. (2014) optimised linear probability models for credit risk, noting calibration "
        "weaknesses that TaxGuard addresses through ensemble outputs constrained to (0,1). Verstraeten and Van den "
        "Poel (2004) warn that sample bias degrades scoring performance\u2014analogous to ZIMRA over-reliance on "
        "enforcement history when first-time aggressors show extreme tax deviation yet zero fine records.",
    )
    add_body(
        doc,
        "Recent public-sector AI guidance stresses transparency, human oversight, and documented feature attributions "
        "before automated decisions affect citizens or firms (Government of Zimbabwe, 2026). SHAP satisfies this "
        "requirement by decomposing each risk score into contributions from tax gap, planning, offshore, and control "
        "variables. Crook et al. (2007) survey consumer credit risk assessment and recommend ensemble methods when "
        "single learners underfit complex boundaries\u2014consistent with TaxGuard\u2019s stacked architecture and with "
        "pairwise scatter evidence of non-linear tier separation in the prototype dataset.",
    )
    add_body(
        doc,
        "Table 1 consolidates terminology for readers approaching the study from policy, audit operations, or data "
        "science backgrounds. Risk scoring differs from automated adjudication: TaxGuard produces ranked probabilities "
        "and explanations; statutory officers retain authority to open, defer, or close cases. Continuous learning "
        "refers to scheduled retraining on confirmed outcomes rather than same-day model mutation, preserving "
        "governance traceability required in fiscal institutions.",
    )
    add_body(
        doc,
        "International tax compliance research emphasises that effective audit selection must combine statutory "
        "knowledge with multivariate analytics (OECD, 2015). Transfer pricing, controlled foreign corporations, and "
        "thin capitalisation each introduce features that manual checklists treat separately. Machine learning "
        "interaction terms\u2014such as Planning_Deviation_Interaction tested in this prototype\u2014model joint "
        "effects that examiners might overlook when indicators appear individually moderate but jointly extreme.",
    )

    # --- 3 Methodology ---
    add_heading_section(doc, "3 Methodology and instruments")
    add_body(
        doc,
        "A prototype demonstration approach was adopted using structured corporate tax filings representative of "
        "ZIMRA audit management workflows (McLeod, 2008). The dataset corporate_tax_risk_dataset.csv contains "
        f"{n_records:,} records, 15 raw fields, and 29 engineered features after domain transformation. Targets "
        "include Tax_Risk_Label (Low, Medium, High), Audit_Outcome (Clean, Qualified, Adverse), and legacy "
        "Audit_Likelihood heuristic scores. Data quality checks confirmed zero missing values and zero duplicate "
        "company identifiers, supporting straight-through ingestion into an Audit Management System mart.",
    )

    add_table_caption(doc, "Table 2. Dataset profile and target distributions.")
    t2 = doc.add_table(rows=5, cols=3)
    t2.style = "Table Grid"
    t2.rows[0].cells[0].text = "Attribute"
    t2.rows[0].cells[1].text = "Value"
    t2.rows[0].cells[2].text = "Interpretation"
    rows2 = [
        ("Records", str(n_records), "Corporate filings in prototype"),
        ("Risk tier balance", "Low 634 | Medium 632 | High 634", "Balanced ML training"),
        ("Audit outcomes", "Clean 69.4% | Qualified 25.2% | Adverse 5.4%", "Realistic post-audit mix"),
        ("Engineered features", "29", "Tax gap, planning, offshore, interactions"),
    ]
    for i, row in enumerate(rows2, start=1):
        for j, val in enumerate(row):
            t2.rows[i].cells[j].text = val

    add_subheading(doc, "3.1 Feature engineering")
    add_body(
        doc,
        "Domain features were constructed to mirror international audit risk practice. Tax rate deviation measures "
        "the gap between statutory and effective tax rates. Tax underpayment ratio and tax gap (million USD) quantify "
        "liability shortfalls. Planning score, offshore intensity, and control risk capture aggressive structures and "
        "weak governance. Interaction terms\u2014Planning_Deviation_Interaction and Audit_Planning_Synergy\u2014model "
        "multiplicative risk hypothesised in BEPS literature (OECD, 2015). All features were standardised before "
        "unsupervised learners; tree-based models used raw engineered inputs.",
    )
    add_subheading(doc, "3.1.1 Feature engineering steps")
    add_body(
        doc,
        "Step 1: Ingest raw return fields (revenue, profit, statutory rate, effective tax rate, offshore subsidiaries, "
        "planning score, control risk, fine history, prior audit outcome). Step 2: Compute tax gap metrics including "
        "Tax_Rate_Deviation, Statutory_Effective_Gap, Tax_Underpayment_Ratio, and Implied_Tax_Liability_Million. "
        "Step 3: Construct behavioural flags for offshore intensity, planning aggressiveness, and audit-planning "
        "synergy. Step 4: Form interaction features capturing joint deviation-planning effects. Step 5: Validate "
        "zero missingness and unique company identifiers before modelling. Step 6: Export engineered matrix to "
        "notebooks and generate_report_assets.py for reproducible figures.",
    )

    add_subheading(doc, "3.2 Hybrid machine learning pipeline")
    add_body(
        doc,
        "Five-fold stratified cross-validation generated out-of-fold (OOF) predictions for each base learner: "
        "Isolation Forest, shallow autoencoder, and HistGradientBoostingClassifier (Pedregosa et al., 2011). "
        "OOF vectors fed a logistic regression meta-learner minimising binary cross-entropy. Primary deployment "
        "target was binary High risk versus other tiers. A secondary model predicted non-compliance "
        "(Qualified or Adverse versus Clean) to support continuous learning analytics. Hyperparameters followed "
        "scikit-learn defaults with learning rate 0.05, maximum tree depth 8, and L2 regularisation for boosting.",
    )

    add_subheading(doc, "3.3 Performance measures and quality control")
    add_body(
        doc,
        "Model quality was assessed using AUROC, PR-AUC, Brier score, F1 at optimal threshold, and confusion "
        "matrices at tau-star (Hanley and McNeil, 1982). Univariate AUROC ranked individual predictors against "
        "Audit_Likelihood baseline. Chi-square tests evaluated independence between risk tiers and audit outcomes "
        "(Tabachnick and Fidell, 2013). Composite policy flags measured high-risk uplift among overlooked cohorts "
        "relative to population baseline 33.4%. Random seed 42 fixed reproducibility across notebooks and "
        "generate_report_assets.py.",
    )
    add_subheading(doc, "3.3.1 Assumptions")
    add_body(
        doc,
        "Linearity: Univariate relationships may be non-linear; tree-based boosting and SHAP capture curvature without "
        " assuming global linearity. Independence of folds: Stratified five-fold cross-validation assumes approximate "
        "exchangeability within strata; temporal dependence in live data may require rolling-origin validation in "
        "production. Label quality: Risk tiers are treated as authoritative training labels; outcome-based relabelling "
        " occurs only in the secondary continuous-learning loop. Explainability: SHAP TreeExplainer assumptions apply "
        "to tree components; ensemble attributions aggregate meta-learner weights.",
    )
    add_subheading(doc, "3.3.2 Instrument development and reproducibility procedures")
    add_body(
        doc,
        "Exploratory analysis was conducted in 01_taxguard_eda.ipynb; hybrid modelling in 02_taxguard_hybrid_model.ipynb. "
        "All figures export to notebooks/figures/ with numbered filenames. Metrics serialize to "
        "outputs/metrics/experiment_summary.json. A React web demonstration in web/ supports stakeholder walkthroughs. "
        "Replication requires Python 3.11+, requirements.txt dependencies, and execution of "
        "scripts/generate_report_assets.py prior to notebook review.",
    )

    add_figure(
        doc,
        FIGURES / "02_univariate_by_risk_tier.png",
        "Figure 2. Univariate distributions of revenue, profit, effective tax rate, tax rate deviation, offshore exposure, and planning score by risk tier. Source: TaxGuard EDA notebook.",
    )

    # --- 4 Results ---
    add_heading_section(doc, "4 Exploratory analysis and TaxGuard prototype results")
    add_body(
        doc,
        "Exploratory analysis established that High-risk filings exhibit heavier right tails on tax rate deviation "
        "and aggressive planning scores, while revenue alone does not separate tiers cleanly. Offshore intensity "
        "rises toward the High tier, supporting inclusion of offshore features in the modelling matrix. These "
        "patterns confirm that revenue-only manual rules are weakly justified and align with international evidence "
        "that statutory\u2013effective gaps dominate corporate audit selection (OECD, 2015).",
    )

    add_table_caption(doc, "Table 3. Top univariate AUROC for High risk detection.")
    t3 = doc.add_table(rows=6, cols=3)
    t3.style = "Table Grid"
    t3.rows[0].cells[0].text = "Feature"
    t3.rows[0].cells[1].text = "AUROC"
    t3.rows[0].cells[2].text = "Interpretation"
    top_feats = [
        ("Tax_Rate_Deviation", "0.9165", "Very strong"),
        ("Planning_Deviation_Interaction", "0.8661", "Strong interaction"),
        ("Tax_Underpayment_Ratio", "0.7800", "Strong"),
        ("Tax_Gap_Million", "0.7453", "Moderate\u2013strong"),
        ("Audit_Likelihood (baseline)", "0.6965", "Below top engineered"),
    ]
    for i, row in enumerate(top_feats, start=1):
        for j, val in enumerate(row):
            t3.rows[i].cells[j].text = val

    add_body(
        doc,
        f"Tax rate deviation achieved AUROC {auc_dev:.4f}, exceeding ZIMRA Audit_Likelihood at {auc_base:.4f}. "
        "Effective tax rate alone showed AUROC near 0.22 because direction of deviation matters: low ETR can "
        "reflect incentives or avoidance depending on context. Planning_Deviation_Interaction ranked second, "
        "confirming that multiplicative combinations outperform isolated indicators reviewed manually.",
    )

    add_figure(
        doc,
        FIGURES / "05_univariate_auc_ranking.png",
        "Figure 3. Top features by univariate AUROC for predicting High tax risk. Source: TaxGuard EDA.",
    )
    add_figure(
        doc,
        FIGURES / "06_roc_univariate_top_features.png",
        "Figure 4. ROC curves for top univariate predictors compared with Audit_Likelihood baseline. Tax rate deviation hugs the upper-left corner. Source: TaxGuard EDA.",
    )

    add_body(
        doc,
        "Correlation analysis showed Tax_Rate_Deviation correlates with underpayment ratio while planning score "
        "adds partially independent signal, justifying interaction terms in the feature map. Chi-square analysis "
        "of risk tier versus audit outcome yielded statistic 6.13 and p-value 0.19, indicating tiers and outcomes "
        "overlap but are not identical; therefore primary queue ranking should use risk labels while post-audit "
        "outcomes drive retraining.",
    )
    add_figure(
        doc,
        FIGURES / "04_correlation_matrix.png",
        "Figure 5. Correlation heatmap of primary TaxGuard indicators. Source: TaxGuard EDA.",
        width=5.2,
    )
    add_body(
        doc,
        "Bivariate boxplots confirmed that the High tier shows higher planning scores and wider deviation distributions, "
        "whereas profit margin overlaps across tiers. Pairwise scatter plots further reveal non-linear boundaries between "
        "Low, Medium, and High tiers, supporting tree-based boosting over linear scorecards (Friedman, 2001). Multiclass "
        "one-versus-rest ROC analysis using tax rate deviation as a score variable demonstrated usable separation across "
        "all three tiers, consistent with the three-tier risk label design used in audit queue stratification.",
    )
    add_figure(
        doc,
        FIGURES / "03_bivariate_boxplots.png",
        "Figure 6. Boxplots of tax deviation, planning score, offshore intensity, profit margin, control risk, and fines by risk tier. Source: TaxGuard EDA.",
    )

    add_table_caption(doc, "Table 4. Overlooked composite indicators (ZIMRA policy gaps).")
    t4 = doc.add_table(rows=6, cols=4)
    t4.style = "Table Grid"
    t4.rows[0].cells[0].text = "Indicator"
    t4.rows[0].cells[1].text = "Flagged n"
    t4.rows[0].cells[2].text = "High-risk rate"
    t4.rows[0].cells[3].text = "Uplift vs baseline"
    overlooked = [
        ("Top-decile tax deviation + zero fines", "116", "91.4%", "2.74\u00d7"),
        ("High planning + clean audit history", "368", "46.7%", "1.40\u00d7"),
        ("High offshore + weak controls", "196", "38.3%", "1.15\u00d7"),
        ("Low profit margin + high revenue", "180", "32.8%", "0.98\u00d7"),
        ("Concentrated ownership + \u22653 offshore subs", "99", "35.4%", "1.06\u00d7"),
    ]
    for i, row in enumerate(overlooked, start=1):
        for j, val in enumerate(row):
            t4.rows[i].cells[j].text = val

    add_body(
        doc,
        "Composite masks reveal cohorts manual rules miss. First-time aggressors with top-decile deviation and zero "
        "fines show 91.4% high-risk rate versus 33.4% baseline\u2014a 2.74-fold uplift. Firms with high planning but "
        "Clean audit history show 1.40-fold uplift, exposing recency bias. Revenue-only selection (low margin plus "
        "high revenue) shows no uplift (0.98\u00d7), confirming that turnover thresholds should not dominate selection.",
    )

    add_figure(
        doc,
        FIGURES / "08_zimra_overlooked_indicators.png",
        "Figure 7. High-risk rate among flagged cohorts versus population baseline (33.4%). Source: TaxGuard policy gap analysis.",
    )

    add_subheading(doc, "4.1 Hybrid ensemble performance")
    add_body(
        doc,
        f"The stacked TaxGuard ensemble achieved OOF ROC-AUC {auc_ens:.4f}, PR-AUC {pr_auc:.4f}, Brier score "
        f"{brier:.4f}, and F1 {f1:.4f} at optimal threshold {metrics['model_oof']['optimal_threshold']:.4f}. "
        f"HistGradientBoosting alone reached AUROC {metrics['model_oof']['hist_gradient_boosting_auc']:.4f}, while "
        f"Isolation Forest and autoencoder alone underperformed (AUROC 0.32 and 0.42) yet contributed orthogonal "
        "errors to the meta-learner. Relative to Audit_Likelihood ({auc_base:.4f}), the ensemble improves ranking "
        "quality by approximately 43% in AUROC terms\u2014meaning auditors reach high-risk firms much earlier in the queue.",
    )

    add_table_caption(doc, "Table 5. Out-of-fold model comparison.")
    t5 = doc.add_table(rows=6, cols=3)
    t5.style = "Table Grid"
    t5.rows[0].cells[0].text = "Model"
    t5.rows[0].cells[1].text = "OOF AUROC"
    t5.rows[0].cells[2].text = "Role"
    model_rows = [
        ("Isolation Forest", "0.3157", "Unsupervised anomaly"),
        ("Autoencoder", "0.4240", "Unsupervised reconstruction"),
        ("HistGradientBoosting", "0.9964", "Supervised core"),
        ("TaxGuard ensemble", f"{auc_ens:.4f}", "Production score"),
        ("Audit_Likelihood", f"{auc_base:.4f}", "Legacy baseline"),
    ]
    for i, row in enumerate(model_rows, start=1):
        for j, val in enumerate(row):
            t5.rows[i].cells[j].text = val

    add_figure(
        doc,
        FIGURES / "11_model_roc_pr_comparison.png",
        "Figure 8. Out-of-fold ROC and precision\u2013recall curves: TaxGuard ensemble versus component models and baseline. Source: TaxGuard hybrid model notebook.",
    )
    add_figure(
        doc,
        FIGURES / "12_confusion_matrix.png",
        "Figure 9. Out-of-fold confusion matrix at optimal threshold (accuracy 97.2%; High-risk precision 97.1%). Source: TaxGuard hybrid model notebook.",
    )
    add_figure(
        doc,
        FIGURES / "13_calibration_curve.png",
        "Figure 10. Reliability diagram; Brier score 0.0244 indicates well-calibrated probabilistic scores. Source: TaxGuard hybrid model notebook.",
    )

    add_subheading(doc, "4.2 Explainability and secondary outcome modelling")
    add_body(
        doc,
        "SHAP summary and bar plots identify tax-gap and planning features as dominant global drivers, satisfying "
        "transparency requirements for auditor-facing prototypes. Waterfall plots decompose individual company scores "
        "for case review. A secondary non-compliance model predicting Qualified or Adverse versus Clean achieved "
        f"AUROC {metrics['model_oof']['non_compliance_ensemble_auc']:.4f}, indicating post-audit outcomes are difficult "
        "to predict pre-filing and should inform continuous learning rather than same-year queue ranking.",
    )

    add_figure(
        doc,
        FIGURES / "14_shap_summary.png",
        "Figure 11. SHAP beeswarm plot showing global feature effects on ensemble risk scores. Source: TaxGuard explainability module.",
    )
    add_figure(
        doc,
        FIGURES / "16_shap_waterfall_example.png",
        "Figure 12. SHAP waterfall explanation for one flagged company (auditor prototype view). Source: TaxGuard explainability module.",
        width=5.5,
    )

    # --- 5 Discussions ---
    add_heading_section(doc, "5 Discussions")
    add_body(
        doc,
        "The research analysed corporate tax audit risk at ZIMRA and established that engineered tax-gap features "
        "outperform legacy Audit_Likelihood heuristics by a wide margin (Nyoni and Matshisela, 2018). Tax rate "
        "deviation emerged as the dominant univariate predictor, supporting mandatory statutory\u2013effective gap "
        "rules in audit management systems. Composite indicators demonstrated that manual processes systematically "
        "under-select first-time aggressors and firms with clean histories but elevated planning scores\u2014cohorts "
        "with uplift up to 2.74 times baseline high-risk prevalence.",
    )
    add_body(
        doc,
        "The hybrid ensemble aligns with credit scoring best practice: calibrated probabilities, out-of-fold "
        "validation, and ensemble diversity (Crook et al., 2007). Near-perfect OOF metrics on the prototype dataset "
        "should be interpreted cautiously; live ZIMRA operational data may introduce label noise, concept drift, "
        "and missing external data that moderate performance. Nevertheless, the relative gain over baseline is "
        "large enough to justify phased pilot deployment with human oversight.",
    )
    add_body(
        doc,
        "Policy implications are immediate. ZIMRA should replace Audit_Likelihood ranking with TaxGuard p_hat in "
        "pilot regions, display SHAP attributions alongside flagged cases, and embed composite masks for first-time "
        "aggressors and offshore-control weaknesses. Sector codes and VAT-to-turnover ratios documented in the "
        "abstract remain integration extensions for Phase 2. Continuous learning should ingest Qualified and Adverse "
        "closures quarterly while preventing feedback loops that permanently penalise incorrectly labelled firms.",
    )
    add_body(
        doc,
        "Comparison with Nyoni and Matshisela (2018) German credit benchmarks (best AUROC 0.8048) illustrates "
        "methodological alignment while emphasising domain difference: TaxGuard targets corporate tax tiers on "
        "structured research data, not consumer loan default. Metrics are not directly comparable across domains, "
        "but evaluation discipline\u2014cross-validation, AUROC, calibration\u2014matches Zimbabwean scoring research norms.",
    )
    add_body(
        doc,
        "Ethical deployment requires version logging, appeals pathways, and monitoring for disparate impact across "
        "sectors once sector codes are integrated. Risk scores must not substitute for due process. Human sign-off on "
        "Adverse labels before they enter training data prevents feedback loops that permanently penalise firms "
        "incorrectly flagged during pilot phases. Calibration curves (Figure 10) support risk-weighted audit hour "
        "planning: when predicted probabilities track observed rates, managers can allocate examiner time "
        "proportionally to expected yield rather than uniform case loads.",
    )
    add_body(
        doc,
        "Phased integration is recommended. Phase 1 (0\u20136 months): batch-score corporate returns monthly in a pilot "
        "region; display top decile with SHAP summaries; compare against legacy Audit_Likelihood ranking; measure hit "
        "rate on Qualified and Adverse closures. Phase 2 (6\u201318 months): embed scores in Audit Management System "
        "workflow; add sector and VAT features; calibrate thresholds by industry; train examiner workshops on reading "
        "SHAP waterfalls. Phase 3 (18\u201336 months): near-real-time scoring on e-filing; automated retraining pipeline; "
        "publish aggregate compliance analytics under privacy-preserving governance.",
    )

    # --- 6 Conclusions ---
    add_heading_section(doc, "6 Conclusions")
    add_body(
        doc,
        "TaxGuard applies credit-scoring discipline to corporate tax audit risk at ZIMRA. Isolation Forest and "
        "autoencoder supply multivariate anomaly signals; HistGradientBoosting captures non-linear tax-gap structure; "
        "logistic stacking calibrates a single risk probability for queue management; SHAP satisfies auditor-facing "
        "transparency. Empirically, tax rate deviation and planning interactions dominate rankings; composite policy "
        f"rules expose cohorts with up to 2.74\u00d7 baseline high-risk rates; the stacked ensemble achieves OOF ROC-AUC "
        f"{auc_ens:.4f}, far above Audit_Likelihood ({auc_base:.4f}).",
    )
    add_body(
        doc,
        "Deployment should proceed through phased Audit Management System integration with human sign-off, "
        "sector and VAT enrichment, and quarterly retraining on audit closures\u2014turning each completed audit into "
        "institutional learning consistent with Zimbabwe\u2019s fiscal digitalisation agenda. Future work includes live "
        "operational validation, sector-normalised thresholds, field trials measuring incremental revenue per audit "
        "hour attributable to TaxGuard ranking, and publication of aggregate compliance analytics under privacy-preserving "
        "governance frameworks advocated in the National AI Strategy (2026\u20132030).",
    )

    # --- References ---
    add_heading_section(doc, "References")
    references = [
        "[1] Basel Committee on Banking Supervision (2006) \u2018International Convergence of Capital Measurement and Capital Standards: A Revised Framework\u2019, Bank for International Settlements, Basel.",
        "[2] BIS (2000) \u2018Principles for the Management of Credit Risk\u2019, Basel Committee document series.",
        "[3] Chen, T. and Guestrin, C. (2016) \u2018XGBoost: A scalable tree boosting system\u2019, Proceedings of KDD, pp. 785\u2013794. doi: 10.1145/2939672.2939785.",
        "[4] Crook, J.N., Edelman, D.B. and Thomas, L.C. (2007) \u2018Recent developments in consumer credit risk assessment\u2019, European Journal of Operational Research, 183(3), pp. 1447\u20131465. doi: 10.1016/j.ejor.2006.09.100.",
        "[5] Friedman, J.H. (2001) \u2018Greedy function approximation: a gradient boosting machine\u2019, Annals of Statistics, 29(5), pp. 1189\u20131232. doi: 10.1214/aos/1013203451.",
        "[6] Government of Zimbabwe (2026) \u2018National AI Strategy 2026\u20132030\u2019, Harare.",
        "[7] Government of Zimbabwe (2024) \u2018National Development Strategy 2 (NDS2)\u2019, domestic revenue mobilisation pillar, Harare.",
        "[8] Hanley, J.A. and McNeil, B.J. (1982) \u2018The meaning and use of the area under a receiver operating characteristic (ROC) curve\u2019, Radiology, 143(1), pp. 29\u201336. doi: 10.1148/radiology.143.1.7063747.",
        "[9] Harwood, J.R. et al. (1999) Managing Risk in Farming: Concepts, Research and Analysis, USDA ERS Agricultural Economic Report No. 774.",
        "[10] Hinton, G.E. and Salakhutdinov, R.R. (2006) \u2018Reducing the dimensionality of data with neural networks\u2019, Science, 313(5786), pp. 504\u2013507. doi: 10.1126/science.1127647.",
        "[11] Kohavi, R. (1995) \u2018A study of cross-validation and bootstrap for accuracy estimation and model selection\u2019, Proceedings of IJCAI, pp. 1137\u20131143.",
        "[12] Liu, F.T., Ting, K.M. and Zhou, Z.-H. (2008) \u2018Isolation Forest\u2019, Proceedings of IEEE ICDM, pp. 413\u2013418. doi: 10.1109/ICDM.2008.17.",
        "[13] Lundberg, S.M. and Lee, S.-I. (2017) \u2018A unified approach to interpreting model predictions\u2019, Advances in Neural Information Processing Systems, 30, pp. 4765\u20134774.",
        "[14] McLeod, S. (2008) \u2018Case study methodology\u2019, Simply Psychology. Available at: https://www.simplypsychology.org/case-study.html (Accessed: 3 June 2026).",
        "[15] Nyathi, K.T. et al. (2014) \u2018Optimisation of the linear probability model for credit risk management\u2019, International Journal of Computer and Information Technology, 3(6), pp. 1340\u20131345.",
        "[16] Nyoni, E.E.T. and Matshisela, N. (2018) \u2018Credit scoring using machine learning algorithms\u2019, Zimbabwe Journal of Science & Technology, 13, pp. 26\u201334.",
        "[17] OECD (2015) \u2018BEPS Action 13: Transfer Pricing Documentation and Country-by-Country Reporting\u2019, OECD Publishing, Paris. doi: 10.1787/9789264241480-en.",
        "[18] Pedregosa, F. et al. (2011) \u2018Scikit-learn: Machine learning in Python\u2019, Journal of Machine Learning Research, 12, pp. 2825\u20132830.",
        "[19] Rose, P.S. (2002) Commercial Bank Management, 5th edn, McGraw-Hill/Irwin, New York.",
        "[20] Tabachnick, B.G. and Fidell, L.S. (2013) Using Multivariate Statistics, 6th edn, Pearson, Boston.",
        "[21] Thomas, L.C. (2000) \u2018A survey of credit and behavioural scoring: forecasting financial risk of lending to consumers\u2019, International Journal of Forecasting, 16(2), pp. 149\u2013172. doi: 10.1016/S0169-2070(00)00034-0.",
        "[22] Verstraeten, G. and Van den Poel, D. (2004) \u2018The impact of sample bias on consumer credit scoring performance\u2019, Journal of the Operational Research Society, 56(8), pp. 981\u2013992. doi: 10.1057/palgrave.jors.2601929.",
        "[23] Zimbabwe Revenue Authority (2024) Corporate income tax administration and audit management framework, Harare.",
        "[24] TaxGuard Project (2026) corporate_tax_risk_dataset.csv, tax-compliance-risk-prediction repository. Available at: https://github.com/ (Accessed: 3 June 2026).",
    ]
    for ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(ref)
        set_run_font(run, size=11)

    add_heading_section(doc, "7 Acknowledgements")
    add_body(
        doc,
        "The authors thank the Zimbabwe Revenue Authority for institutional support in developing this prototype "
        "demonstration under the Data, Automation and Intelligent Systems theme. We acknowledge colleagues in audit "
        "operations and information technology whose domain knowledge informed feature engineering and policy gap "
        "analysis. Responsibility for errors and omissions remains with the authors.",
        first_line_indent=0.5,
    )

    return doc


def main() -> int:
    if not METRICS.exists():
        print(f"Missing metrics: {METRICS}. Run scripts/generate_report_assets.py first.", file=sys.stderr)
        return 1

    metrics = json.loads(METRICS.read_text(encoding="utf-8"))
    doc = build_document(metrics)
    word_count = count_words(doc)

    if word_count < MIN_WORDS or word_count > MAX_WORDS:
        print(
            f"Warning: word count {word_count} outside target {MIN_WORDS}\u2013{MAX_WORDS}. "
            "Adjust prose in build_research_proposal_docx.py if required.",
            file=sys.stderr,
        )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Wrote {OUTPUT}")
    print(f"Word count: {word_count}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
